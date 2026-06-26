import streamlit as st
import os
import sys
import io
from dotenv import load_dotenv
from docx import Document

# Absolute project root directory mapping layout fix for Windows systems
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Load system environment states
load_dotenv()

from src.utils.rag_handler import DocRAGHandler
from src.agents.doc_agents import DocumentationGraphPipeline

def create_word_docx(text_content: str) -> io.BytesIO:
    """Helper function to convert raw text/markdown into a downloadable Word Document byte stream."""
    doc = Document()
    doc.add_heading('Documentation Revision Log', level=1)
    
    # Add the drafted content. For a more advanced setup, you could parse the markdown 
    # strictly into bold/italic docx runs, but raw text injection works perfectly for drafting.
    doc.add_paragraph(text_content)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit App Configurations - Branded for Ivanti
st.set_page_config(
    page_title="Ivanti Documentation Reconciliation Pipeline",
    page_icon="🔴",
    layout="wide"
)

st.title("🔴 Ivanti Technical Documentation Automation Engine")
st.markdown("Reconcile document variations and mitigate content drift by cross-referencing reference manual vectors with engineering pipeline updates.")

# Sidebar Configuration
st.sidebar.header("Pipeline Controls")
gemini_api_key = os.getenv("GEMINI_API_KEY")

if not gemini_api_key:
    st.sidebar.error("⚠️ GEMINI_API_KEY missing from environment configurations (.env file).")
else:
    st.sidebar.success("🔑 Gemini API credentials loaded successfully.")

project_scope = st.sidebar.text_input(
    "Target Engineering Milestone/Sprint Scope",
    value="Sprint 24 - Core Infrastructure Upgrades"
)

st.sidebar.markdown("---")
st.sidebar.info(
    "**System Operational Pipeline:**\n"
    "1. Parse baseline PDF manuals into modular components.\n"
    "2. Vectorize text semantics via Google GenAI Embeddings and ChromaDB.\n"
    "3. Run automated delta processing through a sequential LangGraph chain.\n"
    "4. Output explicit tracking edits directly inside an exportable Markdown file."
)

# Persist operational instances inside global memory structures
if 'rag_handler' not in st.session_state:
    st.session_state.rag_handler = DocRAGHandler()
if 'graph_orchestrator' not in st.session_state:
    st.session_state.graph_orchestrator = DocumentationGraphPipeline()

# User Core Document Asset Drop Zone
uploaded_file = st.file_uploader("Upload Current Product/System Documentation Manual (PDF)", type=["pdf"])

if uploaded_file is not None:
    st.success(f"📄 File '{uploaded_file.name}' successfully loaded into memory buffer.")
    
    # Primary button will automatically pull the Ivanti Red from config.toml
    if st.button("Trigger Documentation Update Pipeline", type="primary"):
        if not gemini_api_key:
            st.error("Please provide valid Gemini keys in your local workspace setup configurations.")
        else:
            with st.spinner("Phase 1: Tokenizing and cataloging reference file layout vectors..."):
                try:
                    total_chunks = st.session_state.rag_handler.ingest_pdf(uploaded_file)
                    st.info(f"✅ Context database built successfully with {total_chunks} distinct reference data blocks.")
                except Exception as e:
                    st.error(f"Failed to process target reference context requirements: {e}")
                    st.stop()

            # Trigger automated execution pipeline
            with st.spinner("Phases 2 & 3: Initializing agent nodes. Constructing structural documentation drafts..."):
                try:
                    pipeline_results = st.session_state.graph_orchestrator.run_pipeline(scope_context=project_scope)
                    
                    st.markdown("---")
                    st.header("📝 Final Drafted Revision Suggestions Output")
                    st.subheader(f"Target Update Baseline Alignment: {project_scope}")
                    
                    # Output markdown updates directly on-screen
                    st.markdown(pipeline_results)
                    st.markdown("---")
                    
                    # Layout formatting: Place download buttons side-by-side
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.download_button(
                            label="Download Markdown Revision Log",
                            data=str(pipeline_results),
                            file_name=f"ivanti_revised_documentation_{project_scope.lower().replace(' ', '_')}.md",
                            mime="text/markdown"
                        )
                        
                    with col2:
                        docx_buffer = create_word_docx(str(pipeline_results))
                        st.download_button(
                            label="Download Word Document (Draft)",
                            data=docx_buffer,
                            file_name=f"ivanti_revised_documentation_{project_scope.lower().replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"An exception was caught during LangGraph data processing loops: {e}")